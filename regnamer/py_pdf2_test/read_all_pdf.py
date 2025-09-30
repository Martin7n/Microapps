from PyPDF2 import PdfReader
from docx import Document


def pdfread(pdf_file_path ):
    reader = PdfReader(pdf_file_path)
    all_text = ""

    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        page = f"\n--- Page {i + 1} ---\n"
        print(f"\n--- Page {i + 1} ---\n")
        print(text if text else "[No extractable text on this page]")
        all_text += text + page if text else ""

    with open("text_extracted.txt", "a", encoding='utf-8') as file:
        file.writelines(all_text)

    return all_text


def pdf_to_word(pdf_file_path, word_output_path):
    reader = PdfReader(pdf_file_path)
    doc = Document()
    doc.add_heading('Extracted PDF Text', 0)

    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        doc.add_paragraph(f"\n--- Page {i + 1} ---\n")
        doc.add_paragraph(text if text else "[No extractable text on this page]")

    doc.save(word_output_path)
    print(f"\n✅ Word file saved to: {word_output_path}")


if __name__ == "__main__":
    # print(pdfread("C:/drob/merged.pdf"))
    pdf_to_word("C:/drob/merged.pdf", "C:/drob/merged_output.docx")
