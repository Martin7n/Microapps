import os
from docx import Document
from PyPDF2 import PdfReader
import re
import openpyxl



import win32com.client as win32
#

## Using built in word...
# word = win32.Dispatch("Word.Application")
# doc = word.Documents.Open(r"C:\path\to\file.doc")
# text = doc.Content.Text
# doc.Close()
# word.Quit()
#
# print(text)



def main_execution(dir_path, phrase):
    file_collected = collect_files(dir_path)
    all_data_content = []
    # all_data_content.extend(doc_reader(file_collected[0]))
    # all_data_content.extend(docx_reader(file_collected[1]))
    all_data_content.extend(pdf_reader(file_collected[2]))
    filtered_content = extract_special_block(all_data_content, phrase)
    #
    # excel_writer(filtered_content, dir_path)


def excel_writer(filtered_content, path_to_dir):
    wb = openpyxl.Workbook()
    sheet = wb.active
    labels = ["Contract","Specials","Name","Size","Type","Direct link"]
    for col, label in enumerate(labels, start=1):
        sheet.cell(row=1, column=col).value = label
    wb.save('extracted_fast_data.xlsx')
    count = 0

    for row, file in enumerate(filtered_content, start=2):
        print(row, file)
        sheet.cell(row=row, column=1).value = file
        sheet.cell(row=row, column=2).value = filtered_content[file]
        count += 1
    wb.save('extracted_fast_data.xlsx')
    wb.close()
    return count

def collect_files(dir_path):
    filelist = os.listdir(dir_path)

    doc_list = []
    docx_list = []
    pdf_list = []
    lst = [doc_list, docx_list,pdf_list]
    for file in filelist:
        filetype = file.split(".")[-1]
        file_key_name = file[0:6]
        if filetype == "docx":
            docx_list.append(os.path.join(dir_path, file))
        elif filetype == "doc":
            doc_list.append(os.path.join(dir_path, file))
        elif filetype=="pdf":
            pdf_list.append(os.path.join(dir_path, file))

    return [doc_list, docx_list, pdf_list]


# def builtIn_reader_doc(file):
#
#     print(file)
#
#     doc = Document()
#     word = win32.Dispatch("Word.Application")
#     doc = word.Documents.Open(file)
#     text = doc.Content.Text
#     doc.Close()
#     word.Quit()
#
#     print(text)


def docx_reader(docx_list):
    result = []
    for file in docx_list:

        doc = Document(file)
        full_text = []


        for para in doc.paragraphs:
            text = para.text.strip()

            if text:
                if len(full_text) == 0:
                    full_text.append(f'{text[0:100]} <tag>')
                full_text.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                full_text.append("\t".join(row_text))

        result.append("\n".join(full_text))
    print(result)
    return result



def pdf_reader(filelist):
    print(filelist)

    result = []

    for file in filelist:
        reader = PdfReader(file)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                if not result:
                    result.append(f'{page_text.strip()[0:100]} <tag>')

                result.append(page_text.strip())

    return result



def extract_special_block(all_data_content, phrase="специални условия", chars_after=300):

    pass
    # resolved = []
    # count = 0
    # # print(all_data_content)
    # print(all_data_content)
    # for text in all_data_content:
    #
    #     match = re.search(re.escape(phrase), text, re.IGNORECASE)
    #     if match:
    #         count += 1
    #         start = match.start()
    #         end = min(start + chars_after, len(text))
    #         key_body = text[start:end]
    #         key_text = text[0:100] + "\n" + str(count)
    #         resolved[key_text] = key_body
    #
    #     print(resolved)

    # return resolved





if __name__ == '__main__':
    dpath = "C:\\Drob"
    phrase = "Special conditions"
    print(main_execution(dpath, phrase))