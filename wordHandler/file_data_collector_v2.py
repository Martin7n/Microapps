import os
from docx import Document
from PyPDF2 import PdfReader
import re
import openpyxl



import win32com.client as win32
from rdflib.plugins.sparql.parserutils import value


#

## Using built in word...
# word = win32.Dispatch("Word.Application")
# doc = word.Documents.Open(r"C:\path\to\file.doc")
# text = doc.Content.Text
# doc.Close()
# word.Quit()
#
# print(text)



def main_execution(dir_path, phrase):
    file_collected = collect_files(dir_path)
    # print(file_collected[1])
    all_data_content = []

    all_found = sum([len(x) for x in file_collected])

    print(f"\033[32m {all_found} discovered... \033[0m  ")
    print(f"\033[32m {len(file_collected[1])} docx to process... \033[0m  ")
    print(f"\033[32m {len(file_collected[2])} pdfs to process... \033[0m  ")

    all_data_content.append(docx_reader(file_collected[1]))
    all_data_content.append(pdf_reader(file_collected[2]))

    filtered_content = extract_special_block(all_data_content, phrase)
    print(f"\033[32m {len(filtered_content)} total to be written... \033[0m  ")

    [rows_written, result_path] = excel_writer(filtered_content, dir_path)

    return f'''\033[32m Total: {rows_written} rows written 
    in {result_path} \n containing <<{phrase}>> phrase \033[0m  '''


def excel_writer(filtered_content, path_to_dir):
    wb = openpyxl.Workbook()
    sheet = wb.active
    labels = ["File","Contract","Specials"]
    for col, label in enumerate(labels, start=1):
        sheet.cell(row=1, column=col).value = label
    wb.save('extracted_dataaaa.xlsx')
    count = 0

    for row, file in enumerate(filtered_content, start=2):
        count += 1
        print(f"\033[32m Writing files...{count}  \033[0m  ")

        sheet.cell(row=row, column=1).value = file
        sheet.cell(row=row, column=2).value = filtered_content[file]
        if "<tag>" in filtered_content[file]:
            header = filtered_content[file].split("<tag>")[0]
        else:
            header = "no tag"
            sheet.cell(row=row, column=1).value = filtered_content[file].split("<tag>")[0]
        sheet.cell(row=row, column=3).value = header

        # sheet.cell(row=row, column=3).value = file["file"]
        # sheet.cell(row=row, column=4).value = file["file"][3]
        # sheet.cell(row=row, column=5).value = file["file"][5]
        # sheet.cell(row=row, column=6).value = file["file"][2]
        # sheet.cell(row=row, column=6).hyperlink = file['file'][2]
        # sheet.cell(row=row, column=6).style = "Hyperlink"
    wb.save('extracted_dataaaa.xlsx')
    wb.close()
    return count, (os.path.join(path_to_dir, 'extracted_dataaaa.xlsx'))

def collect_files(dir_path):
    docs_and_others = [ "doc", "docx" ,"pdf"]
    filelist = os.listdir(dir_path)

    doc_list = []
    docx_list = []
    pdf_list = []
    lst = [doc_list, docx_list,pdf_list]
    for file in filelist:
        filetype = file.split(".")[-1]
        file_key_name = file[0:6]
        if filetype == "docx":
            docx_list.append(os.path.join(dir_path, file))
        elif filetype == "doc":
            doc_list.append(os.path.join(dir_path, file))
        elif filetype=="pdf":
            pdf_list.append(os.path.join(dir_path, file))
    print(pdf_list)
    return [doc_list, docx_list, pdf_list]


# def builtIn_reader_doc(file):
#
#     print(file)
#
#     doc = Document()
#     word = win32.Dispatch("Word.Application")
#     doc = word.Documents.Open(file)
#     text = doc.Content.Text
#     doc.Close()
#     word.Quit()
#
#     print(text)


def docx_reader(docx_list):

    dicty = {}
    for file in docx_list:
        result = []

        doc = Document(file)
        full_text = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                if len(full_text) == 0:
                    tag_text = text[0:100]+"<tag>"
                    full_text.append(tag_text)
                full_text.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                full_text.append("\t".join(row_text))

        result.append("\n".join(full_text))
        dicty[file] = "\n".join(full_text[0:100])

    return dicty



def pdf_reader(filelist):


    dicty = {}

    for file in filelist:
        result = []
        reader = PdfReader(file)

        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                result.append(page_text.strip())
        text_origin = "\n".join(result)
        if text_origin:
            dicty[file] = text_origin

    return dicty



def extract_special_block(all_data_content, phrase="специални условия", chars_after=1500):
    resolved = {}
    print(f"\033[32m Data extraction in progress...  \033[0m  ")
    for records in all_data_content:


        for key, text in records.items():
            match = re.search(re.escape(phrase), text, re.IGNORECASE)
            if match:
                page_title = f'{text[0:50]}<tag>'

                start = match.start()
                end = min(start + chars_after, len(text))
                complete_text = text[start:end]
                partial = complete_text[0:10000]
                # if partial:
                resolved[key] = page_title + partial
            # else:
            #     resolved[key] = "No mached content"

    return resolved


def doc_all_reader(doc_list):

    # Todo: blindly coded...no win32
    dicty = {}
    for file in doc_list:
        full_text = []

        word = win32.Dispatch("Word.Application")
        doc = word.Documents.Open(file)
        text = doc.Content.Text
        doc.Close()

        dicty[file] = "\n".join(full_text[0:100])

    return dicty


def doc_all_reader_2(doc_list):
    dicty = {}

    try:
        from docx import Document as DocxDocument
        docx_available = True
    except ImportError:
        docx_available = False

    try:
        import win32com.client as win32
        word = win32.Dispatch("Word.Application")
        word.Visible = False
        win32_available = True
    except:
        win32_available = False
        word = None

    for file in doc_list:
        ext = os.path.splitext(file)[1].lower()
        text = ""

        if ext == ".docx" and docx_available:
            try:
                doc = DocxDocument(file)
                parts = []

                # paragraphs
                for p in doc.paragraphs:
                    parts.append(p.text)

                # tables
                for table in doc.tables:
                    for row in table.rows:
                        parts.append("\t".join(cell.text for cell in row.cells))

                text = "\n".join(parts)

            except Exception as e:
                text = f"[ERROR reading as DOCX: {e}]"

        elif win32_available:
            try:
                doc = word.Documents.Open(file)
                text = doc.Content.Text
                doc.Close()
            except Exception as e:
                text = f"[ERROR reading via Word: {e}]"

        else:
            text = "[No available method to read file]"

        lines = text.splitlines()
        dicty[file] = "\n".join(lines[:100])

    if word:
        word.Quit()

    return dicty



if __name__ == '__main__':
    dpath = "C:\\Drob"
    phrase = "special conditions"
    print(main_execution(dpath, phrase))