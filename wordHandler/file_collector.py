import os
from docx import Document
from PyPDF2 import PdfReader
import re
import openpyxl



import win32com.client as win32
#

## Using built in word...
# word = win32.Dispatch("Word.Application")
# doc = word.Documents.Open(r"C:\path\to\file.doc")
# text = doc.Content.Text
# doc.Close()
# word.Quit()
#
# print(text)



def main_execution(dir_path):
    file_collected = collect_files(dir_path)
    # print(file_collected[1])
    all_data_content = []
    # all_data_content.extend(doc_reader(file_collected[0]))
    all_data_content.extend(docx_reader(file_collected[1]))
    all_data_content.extend(pdf_reader(file_collected[2]))
    #
    # print(all_data_content[1])
    # print(len(all_data_content))

    # print(all_data_content)
    filtered_content = extract_special_block(all_data_content)

    excel_writer(filtered_content, dir_path)


def excel_writer(filtered_content, path_to_dir):
    wb = openpyxl.Workbook()
    sheet = wb.active
    labels = ["Contract","Specials","Name","Size","Type","Direct link"]
    for col, label in enumerate(labels, start=1):
        sheet.cell(row=1, column=col).value = label
    wb.save('extracted_dataaaa.xlsx')
    count = 0

    for row, file in enumerate(filtered_content, start=2):
        print(row, file)
        sheet.cell(row=row, column=1).value = file
        sheet.cell(row=row, column=2).value = filtered_content[file]
        # sheet.cell(row=row, column=3).value = file["file"]
        # sheet.cell(row=row, column=4).value = file["file"][3]
        # sheet.cell(row=row, column=5).value = file["file"][5]
        # sheet.cell(row=row, column=6).value = file["file"][2]
        # sheet.cell(row=row, column=6).hyperlink = file['file'][2]
        # sheet.cell(row=row, column=6).style = "Hyperlink"
        count += 1
    wb.save('extracted_dataaaa.xlsx')
    wb.close()
    return count

def collect_files(dir_path):
    docs_and_others = [ "doc", "docx" ,"pdf"]
    filelist = os.listdir(dir_path)

    doc_list = []
    docx_list = []
    pdf_list = []
    lst = [doc_list, docx_list,pdf_list]
    for file in filelist:
        filetype = file.split(".")[-1]
        file_key_name = file[0:6]
        if filetype == "docx":
            docx_list.append(os.path.join(dir_path, file))
        elif filetype == "doc":
            doc_list.append(os.path.join(dir_path, file))
        elif filetype=="pdf":
            pdf_list.append(os.path.join(dir_path, file))

    return [doc_list, docx_list, pdf_list]


# def builtIn_reader_doc(file):
#
#     print(file)
#
#     doc = Document()
#     word = win32.Dispatch("Word.Application")
#     doc = word.Documents.Open(file)
#     text = doc.Content.Text
#     doc.Close()
#     word.Quit()
#
#     print(text)


def docx_reader(docx_list):
    result = []
    for file in docx_list:

        doc = Document(file)
        full_text = []


        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)

            # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                full_text.append("\t".join(row_text))

        result.append("\n".join(full_text))

    return result



def pdf_reader(filelist):

    result = []

    for file in filelist:
        # print(file)
        reader = PdfReader(file)
        page = reader.pages[0]
        text_origin = page.extract_text()
        total_length = len(text_origin)
        short = total_length - 50
        text = text_origin[3:300]
        result.append(text_origin)

    return result



def extract_special_block(all_data_content, phrase="специални условия", chars_after=300):
    # phrase = "thermal"
    resolved = {}
    count = 0
    # print(all_data_content)
    for text in all_data_content:

        # key_text = text[0:20]
        key_text = "ZZ"
        key_body = ""

        match = re.search(re.escape(phrase), text, re.IGNORECASE)
        if match:
            count += 1
            start = match.start()
            end = min(start + chars_after, len(text))
            key_body = text[start:end]
            key_text = text[0:50] + "\n" + str(count)
            resolved[key_text] = key_body

    return resolved


def content_collector(content_key, content_body):
    pass

def xls_writer(contents):
    pass



if __name__ == '__main__':
    dpath = "C:\\Drob"
    print(main_execution(dpath))